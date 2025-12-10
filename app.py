# app.py

import io
from flask import Flask, render_template, request, send_file
from docx import Document
# Removed: from docx.shared import Inches (as it was unused)

# --- 1. FLASK APPLICATION SETUP ---
app = Flask(__name__)

# --- Function to generate the DOCX content ---
def create_docx(data):
    """
    Creates the complete DOCX document using all data submitted from the web form.
    Returns the document stream and the final filename.
    """
    doc = Document()
    
    # --- TITLE ---
    doc.add_heading('CHILDREN’S BOOK ILLUSTRATION AGREEMENT', 0)
    # REMOVED: doc.add_paragraph('(Draft for creative practice only — not legal advice)')
    doc.add_paragraph(f'Governing Law: {data.get("governing_law", "Indiana")}')

    # --- PARTIES ---
    doc.add_heading('Parties', level=1)
    doc.add_paragraph('This Agreement is made between:')
    doc.add_paragraph(f'Author/Creator: {data.get("author_name", "[Your Name]")}, residing in {data.get("author_city_state", "[City, State]")}')
    doc.add_paragraph(f'Illustrator (Minor): {data.get("illustrator_name", "[Child’s Name]")}')
    doc.add_paragraph(f'Parent/Legal Guardian: {data.get("guardian_name", "[Parent’s Name]")}, residing in {data.get("guardian_city_state", "[City, State]")}')
    doc.add_paragraph(f'Effective as of: {data.get("effective_date", "[Date]")}')

    # --- SECTIONS (FULL CONTENT) ---
    sections = [
        ("1. Scope of Work",
          f"The Illustrator will create original artwork for a children’s book currently titled \"{data.get('book_title', '[Book Title]')}\". "
          f"Artwork includes: {data.get('deliverables', '[list number/type: cover, page spreads, characters, etc.]')}. "
          f"Deliver files in {data.get('file_format', '[file format]')} by {data.get('deadline', '[deadline]')}. "
          "All drafts and final files will be provided digitally unless otherwise agreed in writing."),
        ("2. Compensation",
          f"Author agrees to pay Illustrator a total of ${data.get('total_compensation', '[amount USD]')} USD as follows:\n"
          "- $_____ upon signing\n"
          "- $_____ upon final delivery of artwork\n\n"
          "Any additional artwork beyond the listed scope must be agreed in writing."),
        ("3. Work-For-Hire & Rights",
          "All artwork is being created as Work-For-Hire under U.S. Copyright law. "
          "Once the Author completes full payment:\n"
          "- All rights, title, and ownership of the artwork transfer to the Author.\n"
          "- The Author receives exclusive worldwide rights to reproduce, publish, modify, distribute, "
          "and commercialize the artwork in any format.\n"
          "- Illustrator and Guardian retain no ongoing control, usage rights, or claim of ownership after payment is complete."),
        ("4. Creative Credit",
          f"In all published versions of the book, Author will include: “Illustrated by {data.get('illustrator_name', '[Child’s Name]') }”. "
          "Author may include biography or artist feature at their discretion."),
        ("5. Confidentiality / NDA",
          "All artwork, story text, drafts, discussions, and creative materials are confidential until the Author releases them publicly. "
          "The Parent/Guardian and Illustrator agree not to share any book content online or with third parties without written permission "
          "from the Author. This confidentiality obligation continues permanently."),
        ("6. Representations",
          "The Guardian confirms:\n"
          "- They are legally authorized to sign on behalf of the minor.\n"
          "- Artwork created is original and does not infringe the rights of others."),
        ("7. Independent Contractor",
          "Illustrator is not an employee, and no employment relationship is created. "
          "No benefits, withholding, or employment rights are provided."),
        ("8. Termination",
          "Either party may terminate this Agreement in writing.\n"
          "If terminated:\n"
          "- Author pays for approved work completed up to that point.\n"
          "- Author keeps rights only to artwork paid for in full.\n"
          "- Illustrator keeps ownership of unpaid artwork; Author may not use it."),
        ("9. Indiana Law",
          "This Agreement is governed by the laws of the State of Indiana."),
        ("10. Entire Agreement",
          "This Agreement represents the full understanding. Any changes must be made in writing and signed by both parties.")
    ]

    for title, text in sections:
        doc.add_heading(title, level=1)
        doc.add_paragraph(text)

    # --- SIGNATURE SECTION ---
    doc.add_heading('Signature Canvas', level=1)

    # Author/Creator
    doc.add_paragraph('Author/Creator')
    p_author = doc.add_paragraph()
    p_author.add_run('Signature: _________________________________').add_break()
    p_author.add_run(f'Printed Name: {data.get("author_name", "________________") }').add_break()
    p_author.add_run(f'Date: {data.get("effective_date", "________________")}')
    
    # Illustrator (Minor)
    doc.add_paragraph('Illustrator (Minor)')
    p_illustrator = doc.add_paragraph()
    p_illustrator.add_run('Signature (optional): ______________________').add_break()
    p_illustrator.add_run(f'Printed Name: {data.get("illustrator_name", "________________") }').add_break()
    p_illustrator.add_run(f'Date: {data.get("effective_date", "________________")}')

    # Parent/Legal Guardian (Required)
    doc.add_paragraph('Parent/Legal Guardian (Required)')
    p_guardian = doc.add_paragraph()
    p_guardian.add_run('Signature: _________________________________').add_break()
    p_guardian.add_run(f'Printed Name: {data.get("guardian_name", "________________") }').add_break()
    p_guardian.add_run(f'Date: {data.get("effective_date", "________________")}')

    # --- INITIALS TABLE ---
    doc.add_heading('Initials', level=2)
    table = doc.add_table(rows=4, cols=3)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Page'
    hdr_cells[1].text = 'Author Initials'
    hdr_cells[2].text = 'Guardian Initials'
    for i in range(1, 4):
        row_cells = table.rows[i].cells
        row_cells[0].text = str(i)
        row_cells[1].text = '______'
        row_cells[2].text = '______'

    # --- EXHIBIT A ---
    doc.add_heading('Exhibit A — Artwork List', level=1)
    exhibit_table = doc.add_table(rows=2, cols=5)
    exhibit_table.style = 'Table Grid'
    headers = ['Illustration', 'Description', 'Format', 'Due Date', 'Status']
    for idx, header in enumerate(headers):
        exhibit_table.rows[0].cells[idx].text = header
    for idx in range(5):
        exhibit_table.rows[1].cells[idx].text = ''

    # --- STREAM THE FILE ---
    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    
    # Generate clean filename
    book_title_safe = data.get('book_title', 'Book').replace(' ', '_').replace('"', '').replace("'", "")
    filename = f"Illustration_Agreement_for_{book_title_safe}.docx"
    
    return file_stream, filename

# --- 2. WEB ROUTES (The two main pages/actions are correct) ---

@app.route('/')
def index():
    return render_template('form.html')
    
@app.route('/generate', methods=['POST'])
def generate():
    form_data = request.form 
    file_stream, filename = create_docx(form_data)
    
    return send_file(
        file_stream, 
        as_attachment=True, 
        download_name=filename,
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )

# --- 3. RUN THE APPLICATION ---

if __name__ == '__main__':
    app.run(debug=True)